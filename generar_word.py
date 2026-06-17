from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────
def set_heading(para, text, level=1):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2e, 0x74, 0xb5)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x5e, 0x8c)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)

def add_h1(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 1)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 2)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 3)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if isinstance(text, list):
        for part, is_bold in text:
            run = p.add_run(part)
            run.font.size = Pt(10.5)
            run.bold = is_bold
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.75)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x1e)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F4F0')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF9E6')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run('⚠  ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x7a, 0x4f, 0x00)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('ONDUPET')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Módulo de Carga y Capacidad')
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x2e, 0x74, 0xb5)

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta_p.add_run('Metodología de Cálculo — Documento Técnico')
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_p.add_run('Versión 1.0  |  Junio 2026')
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
add_h1(doc, '1. Introducción y Alcance')

add_body(doc, (
    'El módulo de Carga y Capacidad de ONDUPET permite visualizar el grado de '
    'ocupación de cada máquina de producción (termoformado y extrusión) para un '
    'horizonte temporal definido por el usuario. El resultado se presenta como un '
    'Mapa de Calor mensual con el porcentaje de saturación de cada recurso.'
))
add_body(doc, (
    'Este documento describe en detalle cómo se calcula la capacidad disponible, '
    'cómo se asigna la carga de las órdenes a cada mes, el criterio de prorrateo '
    'para órdenes de larga duración, y cómo se obtiene el indicador final de '
    'saturación que alimenta el Mapa de Calor.'
))

add_h2(doc, '1.1 Máquinas contempladas')

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
hdrs = ['Código', 'Tipo', 'Turnos por defecto', 'Horas/día (defecto)']
for i, h in enumerate(hdrs):
    shade_cell(tbl.rows[0].cells[i], '1A5376')
    set_cell_text(tbl.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

machines = [
    ('T1', 'Termoformado', '3 turnos', '24 h'),
    ('T2', 'Termoformado', '3 turnos', '24 h'),
    ('T3', 'Termoformado', '3 turnos', '24 h'),
    ('T4', 'Termoformado', '2 turnos', '16 h'),
    ('T5', 'Termoformado', '3 turnos', '24 h'),
    ('T6', 'Termoformado', '2 turnos', '16 h'),
    ('E1', 'Extrusión',    '3 turnos', '24 h'),
    ('E2', 'Extrusión',    '2 turnos', '16 h'),
]
for idx, (cod, tipo, turns, hd) in enumerate(machines):
    row = tbl.add_row()
    fill = 'EBF5FB' if idx % 2 == 0 else 'FDFEFE'
    for c in row.cells:
        shade_cell(c, fill)
    set_cell_text(row.cells[0], cod,   bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], tipo,  bold=False, size=10)
    set_cell_text(row.cells[2], turns, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[3], hd,    bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_body(doc, [
    ('Cada turno equivale a ', False),
    ('8 horas', True),
    (' de producción. Por tanto, una máquina de 3 turnos genera 24 h/día laborable, '
     'una de 2 turnos genera 16 h/día.', False)
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. CAPACIDAD DISPONIBLE
# ══════════════════════════════════════════════════════════════
add_h1(doc, '2. Cálculo de la Capacidad Disponible')

add_body(doc, (
    'La capacidad disponible de una máquina para un mes concreto se obtiene '
    'sumando las horas productivas de cada día del mes. La función principal es:'
))
add_code(doc, 'getCapH(año, mes, máquina, modoCalendario)')
add_body(doc, (
    'Internamente itera sobre todos los días del mes y acumula las horas de cada día:'
))
add_code(doc, 'para cada día d en [1 .. último_día(mes)]:')
add_code(doc, '    horasDía = getDayShiftsCapH(máquina, año, mes, d, modoCalendario)')
add_code(doc, '    capacidadTotal += horasDía')

add_h2(doc, '2.1 Modos de calendario (modoCalendario)')

add_body(doc, (
    'El usuario puede elegir entre los siguientes modos al calcular la carga y capacidad:'
))

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'
for i, h in enumerate(['Modo', 'Días productivos', 'Horas/día']):
    shade_cell(tbl2.rows[0].cells[i], '2E74B5')
    set_cell_text(tbl2.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

modes = [
    ('actual',  'Lunes a viernes, descontando paros registrados', 'Turnos reales de la máquina × 8 h'),
    ('1t',      'Lunes a viernes',                                 '1 turno = 8 h/día'),
    ('2t',      'Lunes a viernes',                                 '2 turnos = 16 h/día'),
    ('3t',      'Lunes a viernes',                                 '3 turnos = 24 h/día'),
    ('4t',      'Lunes a sábado',                                  '3 turnos = 24 h/día'),
    ('5t',      'Todos los días (incluido domingo)',               '3 turnos = 24 h/día'),
    ('custom',  'Definido por el usuario',                         'Turnos configurados por el usuario'),
]
for idx, (m, dias, hrs) in enumerate(modes):
    row = tbl2.add_row()
    fill = 'EAF4FB' if idx % 2 == 0 else 'FDFEFE'
    for c in row.cells:
        shade_cell(c, fill)
    set_cell_text(row.cells[0], m,    bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], dias, bold=False, size=10)
    set_cell_text(row.cells[2], hrs,  bold=False, size=10)

doc.add_paragraph()
add_note(doc, (
    'El modo "actual" es el recomendado para planificación real, ya que descuenta '
    'los paros de mantenimiento, festivos y vacaciones registrados en el calendario.'
))

add_h2(doc, '2.2 Lógica diaria (getDayShiftsCapH)')

add_body(doc, 'El algoritmo para un día concreto sigue esta lógica:')
add_bullet(doc, 'Si hay un paro registrado para esa máquina en esa fecha → 0 horas.')
add_bullet(doc, 'Si el día es fin de semana en el modo aplicable → 0 horas.')
add_bullet(doc, 'En modo "actual": se usan los turnos por defecto de la máquina (dSh) × 8 h.')
add_bullet(doc, 'En modos "1t"–"3t": se usan N turnos × 8 h, ignorando dSh.')
add_bullet(doc, 'En modo "4t"/"5t": 24 h fijas los días que corresponda.')
add_bullet(doc, 'En modo "custom": se usa la configuración de turnos definida por el usuario.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. ASIGNACIÓN DE CARGA — CRITERIO DE FECHA DE INICIO
# ══════════════════════════════════════════════════════════════
add_h1(doc, '3. Asignación de Carga de las Órdenes')

add_h2(doc, '3.1 Criterio general: mes de inicio')

add_body(doc, (
    'Cada orden de fabricación tiene asociada una '
), )
add_body(doc, [
    ('Cada orden de fabricación tiene asociada una ', False),
    ('fecha de inicio (fecha_inicio)', True),
    (' y una ', False),
    ('fecha de fin (fecha_fin)', True),
    (', ademas de las horas totales estimadas (horas_of).', False),
])
add_body(doc, (
    'El criterio general de ONDUPET es: las horas de la orden se asignan al mes '
    'en que comienza la orden, determinado por la fecha_inicio.'
))
add_code(doc, 'mesCarga = fecha_inicio.slice(0, 7)   → "AAAA-MM"')
add_body(doc, (
    'Ejemplo: una orden con fecha_inicio = 2026-06-30 imputa '
    'todas sus horas a junio 2026, aunque la producción continúe en julio.'
))

add_h2(doc, '3.2 Criterio de prorrateo para órdenes de larga duración')

add_note(doc, (
    'Este criterio complementa al anterior y debe aplicarse cuando una '
    'orden abarca más de un mes natural, para evitar que un único mes '
    'absorba todo el impacto y distorsione el indicador de saturación.'
))

add_body(doc, (
    'Cuando una orden comienza en un mes y finaliza en otro (o en meses '
    'posteriores), las horas deben distribuirse proporcionalmente entre '
    'todos los meses que ocupa la orden.'
))

add_h3(doc, '3.2.1 Base del prorrateo: días laborables')

add_body(doc, (
    'La distribución se realiza en base a los días laborables (lunes a '
    'viernes, descontando festivos y paros) que la orden ocupa en cada mes, '
    'no en base a días naturales.'
))
add_body(doc, 'Fórmula general:')
add_code(doc, 'díasLaborables_m  = días laborables de la orden en el mes m')
add_code(doc, 'díasLaborablesTot = suma de días laborables de todos los meses')
add_code(doc, '')
add_code(doc, 'horasEnMes_m = horasTotalesOrden × (díasLaborables_m / díasLaborablesTot)')

add_h3(doc, '3.2.2 Determinación de los días laborables por mes')

add_body(doc, (
    'Para cada mes natural m que la orden ocupa:'
))
add_bullet(doc, (
    'Si es el mes de inicio: se cuentan los días laborables desde '
    'fecha_inicio hasta el último día del mes.'
))
add_bullet(doc, (
    'Si es el mes de fin: se cuentan los días laborables desde el '
    'primer día del mes hasta fecha_fin.'
))
add_bullet(doc, (
    'Si es un mes intermedio: se cuentan todos los días laborables '
    'del mes completo.'
))
add_bullet(doc, (
    'Un día es laborable si es lunes–viernes y no hay paro registrado '
    'para la máquina en esa fecha (coherente con el modo "actual").'
))

add_h3(doc, '3.2.3 Ejemplo ilustrativo')

add_body(doc, [
    ('Supuesto: ', True),
    ('Orden OF-001, máquina T1, 3 turnos (24 h/día laborable).', False),
])
add_bullet(doc, 'fecha_inicio = 2026-06-30 (martes)')
add_bullet(doc, 'fecha_fin    = 2026-07-15 (miércoles)')
add_bullet(doc, 'horas_of     = 360 h')

doc.add_paragraph()
add_body(doc, 'Cálculo de días laborables:')

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = 'Table Grid'
for i, h in enumerate(['Mes', 'Período en la orden', 'Días laborables', 'Horas asignadas']):
    shade_cell(tbl3.rows[0].cells[i], '375E8C')
    set_cell_text(tbl3.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

ex_rows = [
    ('Junio 2026',  '30 jun (1 día)',               '1',    '360 × 1/12 = 30 h'),
    ('Julio 2026',  '1 jul – 15 jul (11 días háb.)', '11',  '360 × 11/12 = 330 h'),
    ('TOTAL',       '—',                              '12',   '360 h'),
]
fills = ['EBF5FB', 'FDFEFE', 'D5E8F5']
for idx, (mes, per, dias, hrs) in enumerate(ex_rows):
    row = tbl3.add_row()
    for c in row.cells:
        shade_cell(c, fills[idx])
    bold_row = (idx == 2)
    set_cell_text(row.cells[0], mes,  bold=bold_row, size=10)
    set_cell_text(row.cells[1], per,  bold=False,    size=10)
    set_cell_text(row.cells[2], dias, bold=bold_row, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[3], hrs,  bold=bold_row, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_body(doc, [
    ('Sin prorrateo, ', False),
    ('junio mostraría 360 h de carga', True),
    (' (saturación potencialmente >100%) aunque en realidad solo se producen 30 h en junio. '
     'Con el prorrateo, el dato refleja la realidad: 30 h en junio y 330 h en julio.', False),
])

add_h3(doc, '3.2.4 Criterio de aplicación del prorrateo')

add_body(doc, 'Se aplica prorrateo cuando se cumple alguna de estas condiciones:')
add_bullet(doc, (
    'La fecha_fin pertenece a un mes diferente a fecha_inicio '
    '(la orden cruza al menos un cambio de mes).'
))
add_bullet(doc, (
    'La duración natural de la orden (fecha_inicio → fecha_fin) supera '
    'los días laborables restantes del mes de inicio.'
))
add_body(doc, 'Si la orden empieza y termina en el mismo mes, no hay prorrateo: todas las horas van al mes de inicio.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. CARGA TOTAL POR MÁQUINA Y MES
# ══════════════════════════════════════════════════════════════
add_h1(doc, '4. Carga Total por Máquina y Mes')

add_h2(doc, '4.1 Estructura de datos de una orden (DCData)')

add_body(doc, (
    'Cada orden en ONDUPET se almacena como un array de 9 elementos:'
))
add_code(doc, 'DCData[máquina][i] = [')
add_code(doc, '    0: código material,')
add_code(doc, '    1: descripción,')
add_code(doc, '    2: estado (Planificado / Previsión / ...),')
add_code(doc, '    3: ym  → "AAAA-MM" (mes al que se imputan las horas),')
add_code(doc, '    4: cantidad (uds),')
add_code(doc, '    5: horas_of (horas totales de la orden),')
add_code(doc, '    6: % rendimiento,')
add_code(doc, '    7: número de OF,')
add_code(doc, '    8: ruta alternativa')
add_code(doc, ']')

add_body(doc, [
    ('El campo ', False),
    ('ym (posición 3)', True),
    (' es el mes de imputación. Actualmente se inicializa con el mes de fecha_inicio. '
     'Con el prorrateo, una orden de larga duración generará múltiples entradas en DCData, '
     'una por cada mes que ocupe, con las horas fraccionadas.', False),
])

add_h2(doc, '4.2 Modos de cálculo: Solo Planificado vs. Plan. + Previsiones')

add_body(doc, (
    'El usuario puede alternar entre dos modos de visualización:'
))

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Modo', 'Qué incluye', 'Uso recomendado']):
    shade_cell(tbl4.rows[0].cells[i], '1A5376')
    set_cell_text(tbl4.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

plan_rows = [
    ('Solo Planificado',
     'Únicamente órdenes con estado "Planificado" (OF confirmadas)',
     'Visión firme del compromiso de capacidad'),
    ('Plan. + Previsiones',
     'Órdenes "Planificado" + órdenes "Previsión" (demanda estimada)',
     'Visión global para detectar cuellos de botella futuros'),
]
for idx, (m, q, u) in enumerate(plan_rows):
    row = tbl4.add_row()
    fill = 'EBF5FB' if idx == 0 else 'FDFEFE'
    for c in row.cells:
        shade_cell(c, fill)
    set_cell_text(row.cells[0], m, bold=True,  size=10)
    set_cell_text(row.cells[1], q, bold=False, size=10)
    set_cell_text(row.cells[2], u, bold=False, size=10)

doc.add_paragraph()

add_h2(doc, '4.3 Acumulación de horas por máquina y mes')

add_body(doc, 'Para cada máquina y cada mes del horizonte:')
add_code(doc, 'cargaMes[máquina][ym] = Σ horas_of  de todas las órdenes')
add_code(doc, '                         con DCData[máq][i][3] == ym')
add_code(doc, '                         y cuyo estado está incluido en el modo activo')

add_h2(doc, '4.4 Tiempo de cambio entre familias (CHG_H)')

add_body(doc, [
    ('Cuando en un mismo mes y máquina hay órdenes de ', False),
    ('diferentes familias de moldes', True),
    (', se añade un tiempo de cambio (changeover) de ', False),
    ('CHG_H = 4 horas', True),
    (' por cada familia adicional presente.', False),
])
add_code(doc, 'familiasMes = número de familias distintas en ese mes/máquina')
add_code(doc, 'horasCambio = (familiasMes - 1) × CHG_H')
add_code(doc, 'cargaMes[máq][ym] += horasCambio')
add_body(doc, (
    'Este tiempo de cambio se refleja en la tabla Detalle por Año-Mes '
    'y se suma a la carga antes de calcular la saturación.'
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. SATURACIÓN
# ══════════════════════════════════════════════════════════════
add_h1(doc, '5. Cálculo de la Saturación (%)')

add_body(doc, 'La saturación de una máquina para un mes es:')

# Formula box
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.space_before = Pt(8)
p_formula.paragraph_format.space_after  = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'EBF5FB')
p_formula._p.get_or_add_pPr().append(shading)
r_f = p_formula.add_run('Saturación (%) = [ Carga (h) / Capacidad disponible (h) ] × 100')
r_f.font.size = Pt(12)
r_f.font.bold = True
r_f.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)

add_body(doc, 'Donde:')
add_bullet(doc, 'Carga (h) = horas imputadas al mes (incluyendo changeovers si aplica).')
add_bullet(doc, 'Capacidad disponible (h) = resultado de getCapH para ese mes y máquina.')

add_body(doc, 'El resultado se redondea al entero más cercano:')
add_code(doc, 'saturación = Math.round(cargaMes / capacidadMes * 100)  [en %]')

add_h2(doc, '5.1 Tratamiento de capacidad cero')

add_body(doc, (
    'Si la capacidad disponible es 0 para un mes (p.ej. la máquina estuvo '
    'parada todo el mes, o el mes no existe en el horizonte), '
    'la saturación se muestra como 0% en lugar de generar un error de división.'
))

add_h2(doc, '5.2 Saturación superior al 100%')

add_body(doc, (
    'Es posible obtener saturaciones superiores al 100% si la carga planificada '
    'supera la capacidad disponible. Esto indica una situación de sobrecarga '
    'que requiere acción: adelantar producción, añadir capacidad, o redistribuir órdenes.'
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. MAPA DE CALOR
# ══════════════════════════════════════════════════════════════
add_h1(doc, '6. Mapa de Calor de Saturación')

add_body(doc, (
    'El Mapa de Calor es la visualización principal del módulo. Presenta una '
    'tabla con máquinas en filas y meses en columnas. Cada celda muestra el '
    'porcentaje de saturación con un código de color que permite identificar '
    'de un vistazo las zonas de riesgo.'
))

add_h2(doc, '6.1 Código de colores')

tbl5 = doc.add_table(rows=1, cols=4)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Color', 'Rango de saturación', 'Significado', 'Acción recomendada']):
    shade_cell(tbl5.rows[0].cells[i], '375E8C')
    set_cell_text(tbl5.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

colors_data = [
    ('d1fae5', '0% – 60%',    'Baja ocupación / máquina disponible', 'Oportunidad para aceptar más trabajo o reducir turnos'),
    ('2dd4a0', '60% – 85%',   'Ocupación óptima',                    'Situación ideal; seguimiento normal'),
    ('f4a261', '85% – 100%',  'Alta ocupación / riesgo leve',         'Revisar fechas de inicio; priorizar órdenes críticas'),
    ('e63946', '> 100%',      'Sobrecarga',                           'Acción inmediata: redistribuir, añadir capacidad o reprogramar'),
]
for idx, (hex_c, rango, sig, acc) in enumerate(colors_data):
    row = tbl5.add_row()
    shade_cell(row.cells[0], hex_c)
    set_cell_text(row.cells[0], '', size=10)
    set_cell_text(row.cells[1], rango, bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], sig,   bold=False, size=10)
    set_cell_text(row.cells[3], acc,   bold=False, size=10)

doc.add_paragraph()

add_h2(doc, '6.2 Estructura del Mapa de Calor')

add_body(doc, 'El Mapa de Calor se organiza en dos pestañas:')
add_bullet(doc, 'Termoformado: máquinas T1, T2, T3, T4, T5, T6.')
add_bullet(doc, 'Extrusión: máquinas E1, E2.')

add_body(doc, (
    'Las columnas representan los meses del horizonte de planificación seleccionado '
    'por el usuario. El encabezado de cada columna muestra el nombre del mes y el '
    'año completo (ej.: "Ene / 2026").'
))
add_body(doc, (
    'Las celdas con saturación 0% se muestran en gris neutro para distinguirlas '
    'de las celdas con baja ocupación (verde claro).'
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. FLUJO COMPLETO DE CÁLCULO
# ══════════════════════════════════════════════════════════════
add_h1(doc, '7. Flujo Completo de Cálculo')

add_body(doc, (
    'A continuación se describe el proceso completo desde que el usuario '
    'lanza el cálculo hasta que ve el Mapa de Calor:'
))

steps = [
    ('1', 'Selección de parámetros',
     'El usuario elige: horizonte temporal (rango de meses), '
     'modo de calendario (actual, 1t–5t, custom) y tipo de plan '
     '(Solo Planificado / Plan.+Previsiones).'),
    ('2', 'Carga de órdenes (DCData)',
     'Se recuperan todas las órdenes de Supabase. Para cada orden: '
     '(a) se asigna al mes de su fecha_inicio; '
     '(b) si la orden es de larga duración (cruza mes), '
     'se prorratea según días laborables por mes.'),
    ('3', 'Acumulación de carga',
     'Para cada combinación (máquina, mes) se suman las horas_of '
     'de las órdenes asignadas a ese mes. Se añaden los tiempos de '
     'cambio CHG_H entre familias distintas.'),
    ('4', 'Cálculo de capacidad',
     'Se llama a getCapH(año, mes, máquina, modoCalendario) para '
     'obtener las horas disponibles de cada máquina en cada mes.'),
    ('5', 'Cálculo de saturación',
     'Saturación (%) = round(carga / capacidad × 100). '
     'Si capacidad = 0, saturación = 0.'),
    ('6', 'Renderizado del Mapa de Calor',
     'Se construye la tabla HTML con los porcentajes y colores correspondientes. '
     'Se actualiza al cambiar de pestaña (Termoformado / Extrusión).'),
]

for num, title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r_num = p.add_run(f'Paso {num}: ')
    r_num.font.size = Pt(11)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)
    r_title = p.add_run(title)
    r_title.font.size = Pt(11)
    r_title.font.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. EJEMPLO COMPLETO
# ══════════════════════════════════════════════════════════════
add_h1(doc, '8. Ejemplo Completo de Cálculo')

add_body(doc, [
    ('Escenario: ', True),
    ('Máquina T1 (3 turnos), junio 2026. Dos órdenes activas.', False),
])

add_h2(doc, '8.1 Capacidad disponible — junio 2026 (T1, modo "actual")')

add_body(doc, (
    'Junio 2026 tiene 30 días naturales. Días laborables (lunes–viernes, '
    'sin festivos ni paros): 22 días.'
))
add_code(doc, 'getCapH(2026, 6, "T1", "actual")')
add_code(doc, '  = 22 días × 3 turnos × 8 h/turno')
add_code(doc, '  = 22 × 24 = 528 h disponibles')

add_h2(doc, '8.2 Órdenes y carga')

tbl6 = doc.add_table(rows=1, cols=6)
tbl6.style = 'Table Grid'
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['OF', 'Inicio', 'Fin', 'Horas tot.', 'Mes(es) imputados', 'Horas en jun-26']):
    shade_cell(tbl6.rows[0].cells[i], '1A5376')
    set_cell_text(tbl6.rows[0].cells[i], h, bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

orders_ex = [
    ('OF-101', '2026-06-10', '2026-06-20', '240 h', 'Jun-26 (solo)',          '240 h (100%)'),
    ('OF-102', '2026-06-30', '2026-07-15', '360 h', 'Jun-26 + Jul-26',        '30 h (8.3%)'),
]
for idx, row_data in enumerate(orders_ex):
    row = tbl6.add_row()
    fill = 'EBF5FB' if idx == 0 else 'FDFEFE'
    for c in row.cells:
        shade_cell(c, fill)
    for ci, val in enumerate(row_data):
        set_cell_text(row.cells[ci], val, bold=False, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_body(doc, 'OF-101 no cruza mes → sin prorrateo. OF-102 sí cruza (30 jun → 15 jul):')
add_code(doc, 'OF-102: días laborables en jun-26 = 1 (solo el 30 jun)')
add_code(doc, '        días laborables en jul-26 = 11 (1 jul–15 jul, lun–vie)')
add_code(doc, '        total días labor. = 12')
add_code(doc, '        horas jun-26 = 360 × (1/12) = 30 h')
add_code(doc, '        horas jul-26 = 360 × (11/12) = 330 h')

add_h2(doc, '8.3 Carga total en junio 2026 (T1)')

add_code(doc, 'carga jun-26 = 240 h (OF-101) + 30 h (OF-102) = 270 h')

add_h2(doc, '8.4 Tiempo de cambio')

add_body(doc, (
    'Supongamos que OF-101 pertenece a familia F01 y OF-102 a familia F02 '
    '→ 2 familias distintas → 1 cambio adicional.'
))
add_code(doc, 'horasCambio = (2-1) × 4 = 4 h')
add_code(doc, 'carga total jun-26 = 270 + 4 = 274 h')

add_h2(doc, '8.5 Saturación')

add_code(doc, 'saturación T1 jun-26 = round(274 / 528 × 100) = round(51.9) = 52%')
add_body(doc, [
    ('Resultado en el Mapa de Calor: ', False),
    ('verde claro (52%, rango 0–60%)', True),
    (' → baja ocupación.', False),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. GLOSARIO
# ══════════════════════════════════════════════════════════════
add_h1(doc, '9. Glosario')

glossary = [
    ('fecha_inicio',     'Fecha en que comienza la fabricación de la orden. Determina el mes base de imputación.'),
    ('fecha_fin',        'Fecha prevista de finalización de la orden.'),
    ('horas_of',         'Horas totales estimadas para completar la orden de fabricación.'),
    ('ym',               'Código de mes de imputación en formato "AAAA-MM". Clave de agrupación en DCData.'),
    ('getCapH',          'Función que calcula las horas disponibles de una máquina en un mes dado.'),
    ('getDayShiftsCapH', 'Función auxiliar que devuelve las horas de un día concreto según el modo de calendario.'),
    ('dSh',              'Diccionario de turnos por defecto de cada máquina: {T1:3, T2:3, T3:3, T4:2, T5:3, T6:2, E1:3, E2:2}.'),
    ('CHG_H',            'Tiempo de cambio de molde entre familias distintas = 4 horas por cambio adicional.'),
    ('DCData',           'Estructura de datos en memoria que almacena todas las entradas de carga por máquina.'),
    ('Saturación (%)',   'Ratio carga/capacidad expresado en porcentaje. Indicador principal del módulo.'),
    ('Prorrateo',        'Distribución proporcional de las horas de una orden entre los meses que ocupa, basada en días laborables.'),
    ('Modo "actual"',    'Calendario real: usa los turnos configurados por máquina y descuenta paros registrados.'),
    ('Modo "Nt"',        'Modo teórico con N turnos fijos (1–5), independiente del calendario de paros.'),
    ('Modo "custom"',    'Calendario personalizado definido por el usuario con turnos específicos por día.'),
    ('Plan. + Prev.',    'Modo que suma órdenes planificadas y previsiones para obtener una visión de demanda total.'),
]

tbl_g = doc.add_table(rows=1, cols=2)
tbl_g.style = 'Table Grid'
tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Término', 'Definición']):
    shade_cell(tbl_g.rows[0].cells[i], '2E74B5')
    set_cell_text(tbl_g.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

for idx, (term, defn) in enumerate(glossary):
    row = tbl_g.add_row()
    fill = 'EAF4FB' if idx % 2 == 0 else 'FDFEFE'
    for c in row.cells:
        shade_cell(c, fill)
    set_cell_text(row.cells[0], term, bold=True,  size=10)
    set_cell_text(row.cells[1], defn, bold=False, size=10)

# ══════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run('ONDUPET · Documento técnico interno · Junio 2026')
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──────────────────────────────────────────────────────
output_path = '/home/user/ONDUPET/ONDUPET_Carga_Capacidad_Metodologia.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
